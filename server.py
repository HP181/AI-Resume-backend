from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import os
import logging
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import io
import json
from fastapi.responses import StreamingResponse

# Create the main app
app = FastAPI()

# Add CORS middleware with more explicit configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://ai-resume-frontend-nu.vercel.app", "*"],  # Add your frontend URL explicitly and allow all origins as fallback
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
    expose_headers=["*"],  # Expose all headers
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Define Models
class ResumeAnalysis(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    extracted_text: str
    missing_sections: List[str]
    weak_areas: List[str]
    improvement_suggestions: List[str]
    improved_resume: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AnalyzeRequest(BaseModel):
    resume_text: str
    target_role: Optional[str] = None

class ExportRequest(BaseModel):
    resume_text: str
    template_style: str  # "modern", "classic", "creative", "professional"
    format: str  # "pdf" or "docx"

# Helper functions for file processing
async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except ImportError as import_error:
        logging.error(f"PyPDF2 import error: {str(import_error)}")
        raise HTTPException(status_code=500, detail="PDF extraction is not available on this server. Please contact support.")
    except Exception as e:
        logging.error(f"PDF extraction error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error extracting PDF text: {str(e)}")

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        import docx
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except ImportError as import_error:
        logging.error(f"python-docx import error: {str(import_error)}")
        raise HTTPException(status_code=500, detail="DOCX extraction is not available on this server. Please contact support.")
    except Exception as e:
        logging.error(f"DOCX extraction error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error extracting DOCX text: {str(e)}")


async def analyze_resume_with_ai(resume_text: str, target_role: Optional[str] = None) -> dict:
    """Analyze resume using OpenAI API with improved error handling and dynamic results"""
    try:
        # Get API key from environment variables
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            logging.error("OPENAI_API_KEY environment variable is missing")
            raise ValueError("OPENAI_API_KEY environment variable is required")
        
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=api_key)
        
        # Create target role context
        target_context = f" for a {target_role} position" if target_role else ""
        
        # Create system message for proper analysis
        system_message = f"""You are an expert resume analyzer and career coach. Analyze resumes{target_context} and provide:
1. Missing sections (e.g., Profile/Summary, Skills, Experience, Education, Certifications, References)
2. Weak areas (vague descriptions, lack of measurable achievements, poor formatting)
3. Specific improvement suggestions with strong action verbs and quantifiable results
4. A polished, improved version of the resume

Be specific and provide personalized feedback based on the content of the resume.
If the resume is technical, focus on technical improvement opportunities.
If the resume is for a non-technical role, focus on relevant skills and achievements for that domain.

Provide your response in this exact JSON format:
{{
  "missing_sections": ["list of missing sections based on actual resume content"],
  "weak_areas": ["specific weak areas identified in this particular resume"],
  "improvement_suggestions": ["detailed, actionable suggestions tailored to this resume"],
  "improved_resume": "complete improved version of the resume with all your suggested enhancements applied"
}}"""

        # Create user message with the resume text
        user_message = f"""Analyze this resume and provide detailed, personalized feedback{target_context}:

{resume_text}

Remember to respond ONLY with valid JSON in the exact format specified."""

        try:
            # Try first with GPT-4 for better analysis if available
            try:
                logging.info("Attempting analysis with GPT-4-turbo")
                response = await client.chat.completions.create(
                    model="gpt-4-turbo-preview",
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ],
                    temperature=0.7,
                    max_tokens=4000
                )
            except Exception as gpt4_error:
                # Fall back to GPT-3.5-turbo if GPT-4 is unavailable
                logging.warning(f"GPT-4 analysis failed: {str(gpt4_error)}. Falling back to GPT-3.5-turbo")
                response = await client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ],
                    temperature=0.7,
                    max_tokens=3000
                )
            
            # Get response text
            response_text = response.choices[0].message.content.strip()
            
            # Clean up response to ensure valid JSON
            if response_text.startswith("```json"):
                response_text = response_text[7:]
            if response_text.startswith("```"):
                response_text = response_text[3:]
            if response_text.endswith("```"):
                response_text = response_text[:-3]
            
            # Parse JSON response
            try:
                analysis_data = json.loads(response_text.strip())
                logging.info("Successfully parsed analysis response from API")
                return analysis_data
            except json.JSONDecodeError as json_err:
                logging.error(f"Failed to parse JSON response: {str(json_err)}")
                logging.error(f"Response text: {response_text[:200]}...")
                # Try with explicit JSON formatting instructions
                return await fallback_analysis(client, resume_text, target_context)
            
        except Exception as api_error:
            logging.error(f"OpenAI API error: {str(api_error)}")
            # Try with a fallback approach
            return await fallback_analysis(client, resume_text, target_context)
            
    except Exception as e:
        logging.error(f"Critical error in AI analysis: {str(e)}")
        logging.warning("*** Using STATIC ANALYSIS as fallback due to critical error ***")
        # Use static analysis as last resort
        return create_static_but_customized_analysis(resume_text)


async def fallback_analysis(client, resume_text, target_context):
    """Fallback analysis using more explicit JSON formatting"""
    try:
        logging.info("Attempting fallback analysis with simplified instructions")
        fallback_system = """You are an expert resume analyzer. You MUST respond ONLY with a valid JSON object containing these exact keys:
- missing_sections: array of strings
- weak_areas: array of strings
- improvement_suggestions: array of strings
- improved_resume: string

Do not include any other text, explanation, or markdown formatting."""

        fallback_prompt = f"""Analyze this resume text and provide specific, personalized feedback{target_context}:

{resume_text}

Respond ONLY with a JSON object."""

        fallback_response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": fallback_system},
                {"role": "user", "content": fallback_prompt}
            ],
            temperature=0.5,
            max_tokens=3000
        )
        
        fallback_text = fallback_response.choices[0].message.content.strip()
        try:
            result = json.loads(fallback_text)
            logging.info("Successfully parsed fallback analysis response")
            return result
        except json.JSONDecodeError:
            logging.error("Failed to parse fallback JSON response")
            logging.warning("*** Using STATIC ANALYSIS as fallback due to JSON parsing error ***")
            return create_static_but_customized_analysis(resume_text)
        
    except Exception as fallback_error:
        logging.error(f"Fallback analysis failed: {str(fallback_error)}")
        logging.warning("*** Using STATIC ANALYSIS as final fallback ***")
        # Only use static response as last resort when everything else fails
        return create_static_but_customized_analysis(resume_text)


def create_static_but_customized_analysis(resume_text):
    """Create a static but somewhat customized analysis based on resume text patterns"""
    logging.warning("*** STATIC ANALYSIS being used - this is not a dynamic AI analysis ***")
    
    # Extract basic sections to see what might be missing
    resume_lower = resume_text.lower()
    
    missing_sections = []
    if not any(term in resume_lower for term in ["summary", "profile", "objective"]):
        missing_sections.append("Professional Summary")
    
    if not any(term in resume_lower for term in ["skill", "expertise", "competenc"]):
        missing_sections.append("Skills Section")
    
    if not any(term in resume_lower for term in ["education", "degree", "university", "college"]):
        missing_sections.append("Education")
    
    if not any(term in resume_lower for term in ["experience", "work", "job", "career"]):
        missing_sections.append("Professional Experience")
    
    # Identify potential weak areas
    weak_areas = []
    if len(resume_text) < 500:
        weak_areas.append("Resume appears too short and lacks detailed content")
    
    if not any(term in resume_lower for term in ["increase", "improve", "achiev", "result", "success"]):
        weak_areas.append("Lack of measurable achievements and results")
    
    # Default weak areas if nothing specific found
    if not weak_areas:
        weak_areas = [
            "Bullet points lack quantifiable achievements",
            "Too much focus on responsibilities rather than accomplishments",
            "No clear career progression highlighted"
        ]
    
    # Generate improvements based on identified issues
    improvements = []
    for area in weak_areas:
        if "achievements" in area.lower():
            improvements.append("Add measurable results to each role (e.g., 'Increased sales by 20%')")
        if "responsibilities" in area.lower():
            improvements.append("Transform duty descriptions into accomplishment statements")
    
    # Default improvements if nothing specific generated
    if not improvements:
        improvements = [
            "Include a skills section with relevant keywords for ATS optimization",
            "Add a professional summary showcasing your unique value proposition",
            "Use strong action verbs at the beginning of bullet points"
        ]
    
    # Add some more specific improvements if we can detect the type of resume
    if "software" in resume_lower or "develop" in resume_lower or "program" in resume_lower:
        improvements.append("Highlight specific programming languages, frameworks, and technical skills")
    elif "market" in resume_lower or "brand" in resume_lower:
        improvements.append("Quantify marketing campaigns with specific ROI and performance metrics")
    elif "sales" in resume_lower:
        improvements.append("Include specific sales figures, quotas achieved, and client acquisition metrics")
    
    logging.warning("Static analysis complete - returning canned response with minimal customization")
    
    return {
        "missing_sections": missing_sections,
        "weak_areas": weak_areas,
        "improvement_suggestions": improvements,
        "improved_resume": resume_text + "\n\n# This would contain an AI-improved version of your resume."
    }

# Debug route to check library availability
@api_router.get("/debug")
async def debug_info():
    """Debug endpoint to check library availability"""
    import sys
    import platform
    
    debug_info = {
        "python_version": sys.version,
        "platform": platform.platform(),
        "libraries": {},
        "environment": {}
    }
    
    # Check for environment variables (redact sensitive values)
    for key in os.environ:
        if key in ['OPENAI_API_KEY', 'MONGO_URL']:
            debug_info["environment"][key] = "Set" if os.environ.get(key) else "Not set"
        else:
            debug_info["environment"][key] = os.environ.get(key)
    
    # Check for library availability
    try:
        import PyPDF2
        debug_info["libraries"]["PyPDF2"] = {
            "installed": True,
            "version": getattr(PyPDF2, "__version__", "unknown")
        }
    except ImportError as e:
        debug_info["libraries"]["PyPDF2"] = {
            "installed": False,
            "error": str(e)
        }
    
    try:
        import docx
        debug_info["libraries"]["python-docx"] = {
            "installed": True,
            "version": getattr(docx, "__version__", "unknown")
        }
    except ImportError as e:
        debug_info["libraries"]["python-docx"] = {
            "installed": False,
            "error": str(e)
        }
    
    try:
        import openai
        debug_info["libraries"]["openai"] = {
            "installed": True,
            "version": getattr(openai, "__version__", "unknown")
        }
    except ImportError as e:
        debug_info["libraries"]["openai"] = {
            "installed": False,
            "error": str(e)
        }
    
    return debug_info

# API Routes
@api_router.get("/")
async def root():
    return {"message": "AI Power Resume API"}

@api_router.options("/{path:path}")
async def options_route(path: str):
    """Handle OPTIONS requests for CORS preflight"""
    return {"detail": "OK"}

@api_router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Upload and extract text from resume"""
    try:
        content = await file.read()
        
        if file.filename.lower().endswith('.pdf'):
            text = await extract_text_from_pdf(content)
        elif file.filename.lower().endswith('.docx'):
            text = await extract_text_from_docx(content)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")
        
        return {
            "success": True,
            "extracted_text": text,
            "filename": file.filename
        }
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error processing file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@api_router.post("/analyze", response_model=ResumeAnalysis)
async def analyze_resume(request: AnalyzeRequest):
    """Analyze resume using AI"""
    try:
        logging.info(f"Starting resume analysis{' for ' + request.target_role if request.target_role else ''}")
        
        analysis_data = await analyze_resume_with_ai(request.resume_text, request.target_role)
        
        analysis = ResumeAnalysis(
            extracted_text=request.resume_text,
            missing_sections=analysis_data.get('missing_sections', []),
            weak_areas=analysis_data.get('weak_areas', []),
            improvement_suggestions=analysis_data.get('improvement_suggestions', []),
            improved_resume=analysis_data.get('improved_resume', '')
        )
        
        # Try to connect to MongoDB if available
        try:
            from motor.motor_asyncio import AsyncIOMotorClient
            
            mongo_url = os.environ.get('MONGO_URL')
            if mongo_url:
                client = AsyncIOMotorClient(
                    mongo_url,
                    maxPoolSize=5,
                    minPoolSize=0,
                    serverSelectionTimeoutMS=5000
                )
                
                db_name = os.environ.get('DB_NAME', 'resume_db')
                db = client[db_name]
                
                doc = analysis.model_dump()
                doc['timestamp'] = doc['timestamp'].isoformat()
                await db.resume_analyses.insert_one(doc)
                logging.info("Analysis saved to database")
            else:
                logging.warning("MongoDB URL not configured, skipping database storage")
        except ImportError:
            logging.warning("MongoDB support not available, skipping database storage")
        except Exception as db_error:
            logging.error(f"Database error: {str(db_error)}")
        
        logging.info("Resume analysis complete and returning results")
        return analysis
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@api_router.post("/export")
async def export_resume(request: ExportRequest):
    """Export resume in specified format and template"""
    try:
        if request.format == "pdf":
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                from reportlab.lib.enums import TA_LEFT, TA_CENTER
                
                buffer = io.BytesIO()
                
                if request.template_style == "professional":
                    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
                    styles = getSampleStyleSheet()
                    
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontSize=20,
                        textColor=colors.HexColor('#1a1a1a'),
                        spaceAfter=12,
                        alignment=TA_CENTER,
                        fontName='Helvetica-Bold'
                    )
                    
                    heading_style = ParagraphStyle(
                        'CustomHeading',
                        parent=styles['Heading2'],
                        fontSize=14,
                        textColor=colors.HexColor('#2c2c2c'),
                        spaceAfter=6,
                        spaceBefore=12,
                        fontName='Helvetica-Bold',
                        borderWidth=1,
                        borderColor=colors.HexColor('#e0e0e0'),
                        borderPadding=5
                    )
                    
                    body_style = ParagraphStyle(
                        'CustomBody',
                        parent=styles['Normal'],
                        fontSize=10,
                        textColor=colors.HexColor('#333333'),
                        spaceAfter=8,
                        fontName='Helvetica'
                    )
                    
                elif request.template_style == "modern":
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    title_style = ParagraphStyle(
                        'ModernTitle',
                        parent=styles['Heading1'],
                        fontSize=22,
                        textColor=colors.HexColor('#0066cc'),
                        spaceAfter=15,
                        alignment=TA_LEFT,
                        fontName='Helvetica-Bold'
                    )
                    
                    heading_style = ParagraphStyle(
                        'ModernHeading',
                        parent=styles['Heading2'],
                        fontSize=13,
                        textColor=colors.HexColor('#0066cc'),
                        spaceAfter=8,
                        spaceBefore=15,
                        fontName='Helvetica-Bold'
                    )
                    
                    body_style = styles['Normal']
                    
                elif request.template_style == "classic":
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    title_style = ParagraphStyle(
                        'ClassicTitle',
                        parent=styles['Heading1'],
                        fontSize=18,
                        textColor=colors.black,
                        spaceAfter=10,
                        alignment=TA_CENTER,
                        fontName='Times-Bold'
                    )
                    
                    heading_style = ParagraphStyle(
                        'ClassicHeading',
                        parent=styles['Heading2'],
                        fontSize=12,
                        textColor=colors.black,
                        spaceAfter=6,
                        spaceBefore=10,
                        fontName='Times-Bold'
                    )
                    
                    body_style = ParagraphStyle(
                        'ClassicBody',
                        parent=styles['Normal'],
                        fontSize=11,
                        fontName='Times-Roman'
                    )
                    
                else:  # creative
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    title_style = ParagraphStyle(
                        'CreativeTitle',
                        parent=styles['Heading1'],
                        fontSize=24,
                        textColor=colors.HexColor('#ff6b35'),
                        spaceAfter=15,
                        alignment=TA_LEFT,
                        fontName='Helvetica-Bold'
                    )
                    
                    heading_style = ParagraphStyle(
                        'CreativeHeading',
                        parent=styles['Heading2'],
                        fontSize=14,
                        textColor=colors.HexColor('#ff6b35'),
                        spaceAfter=8,
                        spaceBefore=12,
                        fontName='Helvetica-Bold'
                    )
                    
                    body_style = styles['Normal']
                
                story = []
                lines = request.resume_text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 0.1*inch))
                        continue
                        
                    if line.isupper() or (len(line) < 50 and line.endswith(':')):
                        story.append(Paragraph(line, heading_style))
                    elif len(story) == 0:
                        story.append(Paragraph(line, title_style))
                    else:
                        story.append(Paragraph(line, body_style))
                
                doc.build(story)
                buffer.seek(0)
                
                # Return streaming response with the generated PDF
                media_type = "application/pdf"
                filename = f"resume_{request.template_style}.pdf"
                
                return StreamingResponse(
                    buffer,
                    media_type=media_type,
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
                
            except ImportError as e:
                logging.warning(f"PDF generation library not available: {str(e)}")
                return {
                    "success": False, 
                    "error": "PDF generation is not available in this deployment. Please try a different format or contact support.",
                    "detail": str(e)
                }
                
        elif request.format == "docx":
            try:
                import docx
                
                doc = docx.Document()
                
                lines = request.resume_text.split('\n')
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue
                        
                    if i == 0:
                        heading = doc.add_heading(line, level=0)
                        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    elif line.isupper() or (len(line) < 50 and line.endswith(':')):
                        doc.add_heading(line, level=1)
                    else:
                        doc.add_paragraph(line)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Return streaming response with the generated DOCX
                media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                filename = f"resume_{request.template_style}.docx"
                
                return StreamingResponse(
                    buffer,
                    media_type=media_type,
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
                
            except ImportError as e:
                logging.warning(f"DOCX generation library not available: {str(e)}")
                return {
                    "success": False, 
                    "error": "DOCX generation is not available in this deployment. Please try a different format or contact support.",
                    "detail": str(e)
                }
                
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")
            
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# For running the app locally
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)